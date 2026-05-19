"""
api/extractor.py — Extracts raw text and structural metadata from PDF and DOCX files.

- extract_text(raw_bytes: bytes, file_type: str) -> dict: Extract text content and metadata
  Returns: {
    raw_text: str,           # all text found
    char_count: int,
    word_count: int,
    has_tables: bool,
    has_images: bool,
    has_columns: bool,
    has_headers_footers: bool,
    page_count: int,
    fonts_used: list[str],
    encoding_issues: list[str],
    extraction_method: str,
    text_blocks: list[dict],  # [{text, x, y, width, height, font_size}] for PDF
    non_text_ratio: float,    # estimated % of page that's non-text (images, graphics)
  }
"""

import io
import logging
import re
from typing import Any

logger = logging.getLogger(__name__)


def extract_text(raw_bytes: bytes, file_type: str) -> dict[str, Any]:
    """Route to correct extractor based on file type."""
    if file_type == "pdf":
        return _extract_pdf(raw_bytes)
    elif file_type in ("docx", "doc"):
        return _extract_docx(raw_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _extract_pdf(raw_bytes: bytes) -> dict[str, Any]:
    """Extract text and metadata from PDF using pdfminer."""
    try:
        import pdfminer.high_level as pdfminer_hl
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import (
            LTPage, LTTextBox, LTTextLine, LTChar, LTFigure,
            LTImage, LTRect, LTLine, LTLayoutContainer
        )
        from pdfminer.pdfpage import PDFPage
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
        from pdfminer.converter import PDFPageAggregator
        from pdfminer.layout import LAParams
    except ImportError:
        raise ValueError("pdfminer.six is required. Install it with: pip install pdfminer.six")

    pdf_file = io.BytesIO(raw_bytes)

    # Basic text extraction
    try:
        raw_text = pdfminer_hl.extract_text(pdf_file)
        pdf_file.seek(0)
    except Exception as e:
        raise ValueError(f"Could not read PDF: {e}")

    if not raw_text:
        raw_text = ""

    # Detailed layout analysis
    pdf_file.seek(0)
    text_blocks = []
    fonts_used = set()
    has_images = False
    has_tables = False
    page_count = 0
    non_text_elements = 0
    total_elements = 0
    encoding_issues = []
    column_x_positions = []

    try:
        laparams = LAParams(
            line_overlap=0.5,
            char_margin=2.0,
            line_margin=0.5,
            word_margin=0.1,
            boxes_flow=0.5,
            detect_vertical=False,
        )
        for page_layout in extract_pages(pdf_file, laparams=laparams):
            page_count += 1
            page_width = page_layout.width
            page_height = page_layout.height

            for element in page_layout:
                total_elements += 1
                if isinstance(element, LTTextBox):
                    block_text = element.get_text().strip()
                    if block_text:
                        x0 = round(element.x0 / page_width * 100, 1)
                        y0 = round((page_height - element.y1) / page_height * 100, 1)
                        width_pct = round((element.x1 - element.x0) / page_width * 100, 1)
                        height_pct = round((element.y1 - element.y0) / page_height * 100, 1)

                        # Collect column positions
                        column_x_positions.append(x0)

                        # Extract font info
                        font_sizes = []
                        for line in element:
                            if isinstance(line, LTTextLine):
                                for char in line:
                                    if isinstance(char, LTChar):
                                        font_name = char.fontname.split("+")[-1] if char.fontname else "Unknown"
                                        fonts_used.add(font_name)
                                        font_sizes.append(round(char.size, 1))

                        avg_font_size = round(sum(font_sizes) / len(font_sizes), 1) if font_sizes else 12.0

                        text_blocks.append({
                            "text": block_text[:500],
                            "x": x0,
                            "y": y0,
                            "width": width_pct,
                            "height": height_pct,
                            "font_size": avg_font_size,
                            "page": page_count,
                        })
                elif isinstance(element, LTFigure):
                    has_images = True
                    non_text_elements += 1
                elif isinstance(element, LTImage):
                    has_images = True
                    non_text_elements += 1
                elif isinstance(element, (LTRect, LTLine)):
                    non_text_elements += 1
    except Exception as e:
        logger.warning(f"Layout analysis partial failure: {e}")

    # Detect multi-column layout
    has_columns = False
    if column_x_positions and len(column_x_positions) > 3:
        left_blocks = sum(1 for x in column_x_positions if x < 20)
        mid_blocks = sum(1 for x in column_x_positions if 40 <= x <= 60)
        right_blocks = sum(1 for x in column_x_positions if x > 60)
        if (left_blocks > 2 and right_blocks > 2) or mid_blocks > 3:
            has_columns = True

    # Detect tables (heuristic: many short text blocks in grid pattern)
    if len(text_blocks) > 10:
        y_positions = [b["y"] for b in text_blocks]
        y_rounded = [round(y / 5) * 5 for y in y_positions]
        from collections import Counter
        y_counts = Counter(y_rounded)
        rows_with_multiple = sum(1 for c in y_counts.values() if c >= 3)
        if rows_with_multiple >= 3:
            has_tables = True

    # Check for encoding issues
    weird_chars = len(re.findall(r'[^\x00-\x7F\u00C0-\u024F\u2000-\u206F]', raw_text))
    replacement_chars = raw_text.count('\ufffd')
    if weird_chars > 10:
        encoding_issues.append(f"{weird_chars} non-standard characters detected")
    if replacement_chars > 0:
        encoding_issues.append(f"{replacement_chars} unreadable characters (replacement chars)")

    # Detect headers/footers (blocks near top/bottom of every page)
    has_headers_footers = False
    if page_count > 1:
        top_blocks = [b for b in text_blocks if b["y"] < 8]
        bottom_blocks = [b for b in text_blocks if b["y"] > 90]
        if len(top_blocks) >= page_count or len(bottom_blocks) >= page_count:
            has_headers_footers = True

    non_text_ratio = round(non_text_elements / max(total_elements, 1), 2)
    words = raw_text.split()

    return {
        "raw_text": raw_text,
        "char_count": len(raw_text),
        "word_count": len(words),
        "has_tables": has_tables,
        "has_images": has_images,
        "has_columns": has_columns,
        "has_headers_footers": has_headers_footers,
        "page_count": page_count,
        "fonts_used": sorted(list(fonts_used))[:10],
        "encoding_issues": encoding_issues,
        "extraction_method": "pdfminer",
        "text_blocks": text_blocks,
        "non_text_ratio": non_text_ratio,
        "file_type": "pdf",
    }


def _extract_docx(raw_bytes: bytes) -> dict[str, Any]:
    """Extract text and metadata from DOCX using python-docx."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        raise ValueError("python-docx is required. Install it with: pip install python-docx")

    try:
        doc = Document(io.BytesIO(raw_bytes))
    except Exception as e:
        raise ValueError(f"Could not read DOCX file: {e}")

    text_parts = []
    text_blocks = []
    fonts_used = set()
    has_tables = False
    has_images = False
    has_columns = False
    encoding_issues = []
    block_index = 0

    # Extract paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            text_parts.append(text)
            # Collect font info
            for run in para.runs:
                if run.font.name:
                    fonts_used.add(run.font.name)
            font_size = 12.0
            for run in para.runs:
                if run.font.size:
                    font_size = round(run.font.size.pt, 1)
                    break

            text_blocks.append({
                "text": text[:500],
                "x": 5.0,
                "y": round(block_index / max(len(doc.paragraphs), 1) * 90 + 5, 1),
                "width": 90.0,
                "height": 3.0,
                "font_size": font_size,
                "page": 1,
            })
            block_index += 1

    # Extract table text
    for table in doc.tables:
        has_tables = True
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                text_parts.append(" | ".join(row_texts))

    # Check for images
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            has_images = True
            break

    # Check for multi-column sections
    for section in doc.sections:
        cols = section._sectPr.findall(qn('w:cols'))
        for col_el in cols:
            num = col_el.get(qn('w:num'))
            if num and int(num) > 1:
                has_columns = True

    raw_text = "\n".join(text_parts)
    words = raw_text.split()

    # Estimate page count
    page_count = max(1, len(words) // 350)

    return {
        "raw_text": raw_text,
        "char_count": len(raw_text),
        "word_count": len(words),
        "has_tables": has_tables,
        "has_images": has_images,
        "has_columns": has_columns,
        "has_headers_footers": False,
        "page_count": page_count,
        "fonts_used": sorted(list(fonts_used))[:10],
        "encoding_issues": encoding_issues,
        "extraction_method": "python-docx",
        "text_blocks": text_blocks,
        "non_text_ratio": 0.05 if has_images else 0.0,
        "file_type": "docx",
    }
